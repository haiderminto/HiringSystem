import os
import logging

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_path: str) -> str:
    """Extract plain text from a DOCX file using python-docx."""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs).strip()
    except Exception as e:
        logger.error(f"DOCX text extraction failed: {e}")
        return ""


def convert_docx_to_pdf(docx_path: str) -> dict:
    """
    Convert a DOCX file to PDF using docx2pdf.

    Returns:
        dict with keys: success, pdf_path, error
    """
    pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        if os.path.exists(pdf_path):
            return {"success": True, "pdf_path": pdf_path, "error": None}
        else:
            return {"success": False, "pdf_path": None, "error": "PDF file not created after conversion"}
    except ImportError:
        logger.warning("docx2pdf not available, falling back to text-only mode")
        return {"success": False, "pdf_path": None, "error": "docx2pdf not installed"}
    except Exception as e:
        logger.error(f"DOCX to PDF conversion failed: {e}")
        return {"success": False, "pdf_path": None, "error": str(e)}


def validate_docx(file_path: str) -> dict:
    """
    Validate a DOCX file.

    Returns:
        dict with keys: valid, error
    """
    try:
        from docx import Document
        doc = Document(file_path)
        _ = len(doc.paragraphs)
        return {"valid": True, "error": None}
    except Exception as e:
        logger.error(f"DOCX validation failed: {e}")
        return {"valid": False, "error": str(e)}
